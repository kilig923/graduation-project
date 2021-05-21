from report_sentence_generation2 import remakeSentence
from docx import Document
import runSql as rs
from docx.shared import Pt, Cm  # 磅数
from coverAndPageMarkAnd2PDF import getCoverPage, addPageMark, doc2PDF, \
    chg_font, delete_paragraph, addModuleTitle1, addModuleTitle2
from docxcompose.composer import Composer
import os


# 从数据库中获取图表
def getimage(sentence_id):
    try:
        sql = "select * from macro_child_report_graph_date where moduleId = "+str(sentence_id)
        result = rs.SuccessSql(sql)[0]
        img_dir_path = "./图片/"+str(sentence_id)
        if not os.path.exists(img_dir_path):  # 文件夹不存在 创建文件夹
            os.makedirs(img_dir_path)
        img_path = os.path.join(img_dir_path, "%s.jpg" % result["id"])
        # 存下图片
        f = open(img_path, 'wb')
        f.write(result["image"])
        return img_path
    except Exception as e:
        import traceback
        traceback.print_exc()
        return ""

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 导入段落对齐包


def remakeParagraphLevelContent(reportName, doc,date,area):
    sql = "select * from macro_child_report_module_indicator where moduleName = " \
          "'"+str(reportName)+"' and isDelete = 0 and area='"+area+"'"
    result = rs.SuccessSql(sql, isSelect=True)
#    print(result)
    module_content = ""
    if(len(date)==4):
        enddate=str(date)+'-12-31'
    else:
        enddate=str(date)+'-31'
    for num in range(len(result)):
        print(result[num])
        indicatorId=result[num]['indicatorId']

        sql = "select * from macro_child_report_module_indicator_date where " \
              "indicatorId = '%s' and date<='%s' and area='%s' and isDelete = 0  order by date desc"\
              %(indicatorId,enddate,area)
        res = rs.SuccessSql(sql, isSelect=True)[0]

        if res['result'] is not None and res['result'] != "":
            paragraph = res['result'].lstrip(' ').replace('\n','')
        else:
            paragraph = ""

        module_content = module_content + paragraph
        if paragraph != "":  # 本句有内容
            # word：放入内容、设置格式
            paragraph1 = doc.add_paragraph()
            ph_format = paragraph1.paragraph_format
            ph_format.space_before = Pt(10)  # 设置段前间距
            ph_format.space_after = Pt(30)  # 设置段后间距
            ph_format.line_spacing = Pt(25)  # 设置行间距
            ph_format.first_line_indent = 406400  # 行开头缩进2字符
            run = paragraph1.add_run(paragraph)  # 创建run，另设字号、粗细
            chg_font(run, fontname='微软雅黑', size=Pt(12))  # 设置字体和字号
            # 添加图片
            img_path = getimage(res['id'])
            if img_path != "":  # 能获取到图片
                paragraph2 = doc.add_paragraph()
                run2 = paragraph2.add_run()  # 创建run，另设字号、粗细
                inline_shape = run2.add_picture(img_path)  # 放入图片
                scale = inline_shape.height / inline_shape.width  # 计算图片长/宽比例
                inline_shape.width = Cm(16)  # 设置宽度
                inline_shape.height = int(scale * inline_shape.width)  # 计算新高度 必须是int
                paragraph2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中

    content = module_content.strip('\n')
    if content != "":
        return module_content, doc
    else:
        return "", doc  # 该段落无内容










def remakeParagraphLevel2Content_2(title, doc):
    sql = "select * from macro_child_report_module_indicator where moduleName = '"+\
          str(title)+"' and isDelete = 0"
    result = rs.SuccessSql(sql, isSelect=True)
   # print(result)
    module_content = ""
    for num, res in enumerate(result):
        # 获取该句/段内容 ---------------------------
        # paragraph = remakeSentence(res['id']).replace('\n', '')
        # --------------------------- 获取该句/段内容

        # 测试快速获取以前生成的句子组成文档 ---------------------------
        if res['result'] is not None and res['result'] != "":
            paragraph = res['result'].split(' ').replace('/n','')
        else:
            paragraph = ""
        # --------------------------- 测试快速获取以前生成的句子组成文档

        module_content = module_content + paragraph
        if paragraph != "":  # 本句有内容
            # word：放入内容、设置格式
            paragraph1 = doc.add_paragraph()
            ph_format = paragraph1.paragraph_format
            ph_format.space_before = Pt(10)  # 设置段前间距
            ph_format.space_after = Pt(12)  # 设置段后间距
            ph_format.line_spacing = Pt(25)  # 设置行间距
    #        ph_format.first_line_indent = 406400  # 行开头缩进2字符
            run = paragraph1.add_run(paragraph)  # 创建run，另设字号、粗细
            chg_font(run, fontname='微软雅黑', size=Pt(12))  # 设置字体和字号
            # 添加图片
            img_path = getimage(res['id'])
            if img_path != "":  # 能获取到图片
                paragraph2 = doc.add_paragraph()
                run2 = paragraph2.add_run()  # 创建run，另设字号、粗细
                inline_shape = run2.add_picture(img_path)  # 放入图片
                scale = inline_shape.height / inline_shape.width  # 计算图片长/宽比例
                inline_shape.width = Cm(16)  # 设置宽度
                inline_shape.height = int(scale * inline_shape.width)  # 计算新高度 必须是int
                paragraph2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中

    content = module_content.strip('\n')
    if content != "":
        return module_content, doc
    else:
        return "", doc  # 该段落无内容


# def remakeParagraphLevel2Content(moduleId, doc):
#     sql = "select id from macro_report_module_indicator where moduleId = "+str(moduleId)+" and isDelete = 0"
#     result = rs.SuccessSql(sql, isSelect=True)
#     paragraph = ""
#     for num, res in enumerate(result):
#         # 获取该句/段内容
#         sentence = remakeSentence(res['id']).replace('\n', '')
#         paragraph = paragraph + sentence
#
#     if paragraph != "":
#         # word：放入内容、设置格式
#         paragraph1 = doc.add_paragraph()
#         ph_format = paragraph1.paragraph_format
#         ph_format.space_before = Pt(10)  # 设置段前间距
#         ph_format.space_after = Pt(12)  # 设置段后间距
#         ph_format.line_spacing = Pt(25)  # 设置行间距
#         ph_format.first_line_indent = 406400  # 行开头缩进2字符
#         run = paragraph1.add_run(paragraph)  # 创建run，另设字号、粗细
#         chg_font(run, fontname='微软雅黑', size=Pt(12))  # 设置字体和字号
#     para = paragraph.strip('\n')
#     if para != "":
#         return paragraph, doc
#     else:
#         return "", doc  # 该段落无内容


# def remakeParagraphLevel1Content(parentId, doc):
#     paragraph = ""  # 段落文本
#
#     sql = "select id, subModuleName from macro_report_module_level2 where parentId = "+str(parentId)+" and isDelete = 0"
#     result = rs.SuccessSql(sql, isSelect=True)
#     for res in result:
#         # 添加二级标题 （可能该模块会没有内容，后续可能将此标题（最后一段）删除）
#         doc = addModuleTitle2(doc, res['subModuleName'])
#
#         paragraph2, doc = remakeParagraphLevel2Content_2(res['id'], doc)
#         if paragraph2.strip('\n') is "":  # 段落内容为空，word要删除本模块标题（此时的最后一段），pdf不做添加即可
#             delete_paragraph(doc.paragraphs[len(doc.paragraphs)-1])
#         else:
#             paragraph = paragraph + paragraph2  # 文本
#     return paragraph, doc


def remakeOneReport(reportId):
    print("[", str(reportId), "START ]")
    sql = "select * from macro_child_report_category_date where id = "+str(reportId)
#    sql = "select * from macro_child_report_category_date where id = "+str(reportId)+" and isDelete = 0"
    results=rs.SuccessSql(sql, isSelect=True)[0]
    reportName=results['reportName']
    parentId=results['reportId']
    date=results['date']
    area=results['area']
    # 标题摘要部分 =================================================================================
    BigreportName = reportName.split('_')[0]
    indexFreq = reportName.split('_')[1]
    if(area!='-1'):
        if(len(area)==6):
            sql = "select area_name from area where code = " + str(area)
            res1 = rs.SuccessSql(sql, isSelect=True)[0]
            areaname=res1['area_name']
        else:
            areaname=area
    else:
        areaname=''
    if(indexFreq=='年'):
        reportType='年度分析报告'
        timename=date+'年'
    elif(indexFreq=='季'):
        reportType='季度分析报告'
        year=date.split('-')[0]
        month=date.split('-')[1]
        if(month=='3' or month=='03'):
            timename=year+'年第一季度'
        elif(month=='6' or month=='06'):
            timename=year+'年第二季度'
        elif(month=='9' or month=='09'):
            timename=year+'年第三季度'
        elif(month=='12' or month=='12'):
            timename=year+'年第四季度'
        else:
            timename=''
    elif(indexFreq=='月'):
        reportType='月度分析报告'
        year = date.split('-')[0]
        month = date.split('-')[1]
        if (month == '1' or month == '01'):
            timename = year + '年一月份'
        elif (month == '2' or month == '02'):
            timename = year + '年二月份'
        elif (month == '3' or month == '03'):
            timename = year + '年三月份'
        elif (month == '4' or month == '04'):
            timename = year + '年四月份'
        elif (month == '5' or month == '05'):
            timename = year + '年五月份'
        elif (month == '6' or month == '06'):
            timename = year + '年六月份'
        elif (month == '7' or month == '07'):
            timename = year + '年七月份'
        elif (month == '8' or month == '08'):
            timename = year + '年八月份'
        elif (month == '9' or month == '09'):
            timename = year + '年九月份'
        elif (month == '10' or month == '10'):
            timename = year + '年十月份'
        elif (month == '11' or month == '11'):
            timename = year + '年十一月份'
        elif (month == '12' or month == '12'):
            timename = year + '年十二月份'
        else:
            timename=''
    if(area=='-1'):
        title=timename+'关于'+BigreportName+'的数据分析报告'
        abstract = "截至" + date + "年末，主要针对"+BigreportName\
                   +"指标的数据进行分析和处理，生成相应的"+reportType
    else:
        BigreportName= BigreportName.lstrip('各省').lstrip('各市')
        title = timename +areaname+ '关于' + BigreportName + '的数据分析报告'
        abstract = "截至" + date + "年末，主要针对" +areaname+"的"+ BigreportName \
                   + "指标的数据进行分析和处理，生成相应的" + reportType
    print(abstract,title)

    # 制作封面文件 =================================================================================
    print("制作封面...")
    path=os.getcwd()
    coverPage_filepath, doc \
        = getCoverPage(path+"/报告/",
                       path+"/报告/"+str(reportId),
                       title, "", "云报告工作室")  # —— 副标题十八字副标题十八字副标题十八字
    doc.add_page_break()  # 添加分页


    # 正文部分 =================================================================================
    print("添加正文...")
    paragraph2, doc = remakeParagraphLevelContent(reportName, doc,date,area)

    # paragraph, doc = remakeParagraphLevel1Content(res['id'], doc)
    # sql = "select id from macro_report_module_level1 where reportId = "+str(reportId)+" and isDelete = 0"  # limit 0, 1
    # result = rs.SuccessSql(sql, isSelect=True)
    # for res in result:
    #     # 添加一级标题 （可能该模块会没有内容，后续可能将此标题（最后一段）删除）
    #    # doc = addModuleTitle1(doc, res['moduleName'])
    #
    #     # 获取该模块内容
    #     paragraph, doc = remakeParagraphLevel1Content(res['id'], doc)
    #     if paragraph.strip('\n') == "":  # word要删除本模块标题（此时的最后一段）
    #         delete_paragraph(doc.paragraphs[len(doc.paragraphs)-1])
    doc.add_page_break()  # 添加分页

    # 添加尾页 =================================================================================
    print("添加尾页...")
    composer = Composer(doc)
    lastpage_filepath = "./报告/lastpage_zyb_纯生成报告用.docx"
    doc2 = Document(lastpage_filepath)
    composer.append(doc2)
    output_all_filepath = "./报告/"+str(reportId)+"/"+title+"_封面_正文_尾页.docx"
    composer.save(output_all_filepath)  # 中间文件 后面删去

    # 添加页码 word文件生成完毕 =================================================================================
    print("添加页码...")
    output_filepath = "./报告/"+str(reportId)+"/"+title+".docx"
    addPageMark(output_all_filepath, output_filepath)
    print("word文件生成完毕，删去中间文件...")
    os.remove(output_all_filepath)  # 删去中间文件
    os.remove(coverPage_filepath)  # 删去封面word文件

    # PDF =================================================================================
    print("word转PDF...")
    pdf_file = output_filepath.replace(".docx", ".pdf")
    # 【水印 待改成“智研报”】
    output_filepath = path + output_filepath
    output_filepath = output_filepath.replace('./', '/').replace('\\', '/')
    pdf_file = path + pdf_file
    pdf_file = pdf_file.replace('./', '/').replace('\\', '/')
    print(output_filepath)
    print(pdf_file)
    doc2PDF(output_filepath, pdf_file)
    print("[", str(reportId), "FINISHED ]")
    return title, output_filepath, pdf_file,date,abstract  # 返回word地址、pdf地址


def remakeReport():
    # sql = "select id from macro_child_report_category"  # limit 0, 1
    # result = rs.SuccessSql(sql)
    # for res in result:
    #28549,28600 含有省份的数据
    #245672, 245673  area对应位areacode
    #517585,517586 area对应的指标为code 未更新
    for i in range(517585,517586):
        # 生成文本报告，返回文件地址、标题
        title, wordUrl, pdfUrl,date,abstract = remakeOneReport(i)#res['id'])
        #插入数据库，新的报告
        sql = "insert into report_file_produce_date(title, source, path, wordPath, wordCopyState," \
              " success, isDelete,publish_date,description) " \
              "values('"+title+"', 'self', '"+pdfUrl+"', '"+wordUrl+"', 0, 0, 0,'"+date+"','"+abstract+"')"
        print(sql)
        # rs.RunSql_data(sql, isSelect=False)


remakeReport()


