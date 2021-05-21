import time
from docx import Document
from docx.shared import Pt, RGBColor, Cm  # 磅数
from docx.oxml.ns import qn  # 中文格式
import traceback
from win32com import client
import os


# ========================================== word正文生成相关 =========================================


# word生成相关字体设置
def chg_font(obj, fontname='微软雅黑', size=None):
    # 设置字体函数
    obj.font.name = fontname
    obj._element.rPr.rFonts.set(qn('w:eastAsia'), fontname)
    if size and isinstance(size, Pt):
        obj.font.size = size


# 删除word段落
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    # p._p = p._element = None
    paragraph._p = paragraph._element = None


# 一级标题
def addModuleTitle1(doc, content):
    paragraph1 = doc.add_paragraph()  # 创建段落，此处也可直接放入文本，默认字号11
    ph_format = paragraph1.paragraph_format
    ph_format.space_before = Pt(10)  # 设置段前间距
    ph_format.space_after = Pt(12)  # 设置段后间距
    ph_format.line_spacing = Pt(20)  # 设置行间距
    run = paragraph1.add_run(content)  # 创建run，另设字号、粗细
    run.bold = True  # 设置字体为粗体
    chg_font(run, fontname='微软雅黑', size=Pt(14))  # 设置字体和字号
    return doc


# 二级标题
def addModuleTitle2(doc, content):
    paragraph1 = doc.add_paragraph()  # 创建段落，此处也可直接放入文本，默认字号11
    ph_format = paragraph1.paragraph_format
    ph_format.space_before = Pt(10)  # 设置段前间距
    ph_format.space_after = Pt(12)  # 设置段后间距
    ph_format.line_spacing = Pt(19)  # 设置行间距
    ph_format.first_line_indent = 406400  # 行开头缩进2字符
    run = paragraph1.add_run(content)  # 创建run，另设字号、粗细
    run.bold = True  # 设置字体为粗体
    chg_font(run, fontname='微软雅黑', size=Pt(13))  # 设置字体和字号
    return doc


# 生成初始报告模块，设置为word一般设置，包括页边距等
def createNewDocument(doc_filepath):
    if doc_filepath != "":
        doc = Document(doc_filepath)  # 以已有模板建立文档对象
    else:
        doc = Document()  # 以默认模板建立文档对象
    sec = doc.sections[0]  # sections对应文档中的“节”
    # 以下依次设置左、右、上、下页面边距
    # 改用使用“适中”页面布局页边距
    distance1 = Cm(2.54)
    distance2 = Cm(1.91)
    sec.left_margin = distance2
    sec.right_margin = distance2
    sec.top_margin = distance1
    sec.bottom_margin = distance1
    # 设置页面宽度 高度
    sec.page_width = Cm(20.9)
    sec.page_height = Cm(29.6)
    # 设置默认字体
    chg_font(doc.styles['Normal'], fontname='宋体')
    return doc


# ========================================== word封面生成相关 =========================================


def addEmptyRows(doc, number_rows, distance):
    row = ""
    for i in range(0, number_rows):
        row += "\n"
    paragraph0 = doc.add_paragraph(row)  # 创建段落，此处也可直接放入文本，默认字号11
    ph_format = paragraph0.paragraph_format
    ph_format.space_before = Pt(10)  # 设置段前间距
    ph_format.space_after = Pt(10)  # 设置段后间距
    ph_format.line_spacing = Pt(distance)  # 设置行间距
    return doc


def addTitle(doc, title):
    # 对标题大小做一定适配：8以内48，8-15字对应48-36，15以上默认36
    if len(title) <= 8:
        font_size = 48
    elif len(title) <= 15:
        font_size = int(36+(15-len(title)))
    else:
        font_size = 36
    # header = doc.sections[0].headers(client.constants.wdHeaderFooterPrimary).Range
    # header.Text = title
    # header.Paragraphs.Alignment = client.constants.wdAlignParagraphCenter
    paragraph1 = doc.add_paragraph()  # 创建段落，此处也可直接放入文本，默认字号11
    ph_format = paragraph1.paragraph_format
    ph_format.space_before = Pt(10)  # 设置段前间距
    ph_format.space_after = Pt(12)  # 设置段后间距
    ph_format.line_spacing = Pt(50)  # 设置行间距
    paragraph1.alignment = 2  # client.constants.wdAlignParagraphCenter 居右
    run = paragraph1.add_run(title)  # 创建run，另设字号、粗细
    run.bold = True  # 设置字体为粗体
    run.font.color.rgb = RGBColor(240, 220, 220)
    chg_font(run, fontname='微软雅黑', size=Pt(font_size))  # 设置字体和字号
    num_empty_rows = 8 - int((len(title)-1)/12)
    return doc, num_empty_rows


def addSubTitle(doc, subTitle, font_size):
    paragraph1 = doc.add_paragraph()  # 创建段落，此处也可直接放入文本，默认字号11
    ph_format = paragraph1.paragraph_format
    ph_format.space_before = Pt(10)  # 设置段前间距
    ph_format.space_after = Pt(12)  # 设置段后间距
    ph_format.line_spacing = Pt(30)  # 设置行间距
    ph_format.first_line_indent = 203200*2  # 行开头缩进2字符
    run = paragraph1.add_run(subTitle)  # 创建run，另设字号、粗细
    run.bold = True  # 设置字体为粗体
    run.font.color.rgb = RGBColor(40, 40, 40)
    chg_font(run, fontname='微软雅黑', size=Pt(font_size))  # 设置字体和字号
    return doc


def addRightStr(doc, nameStr, number_line_indent, font_size):
    paragraph1 = doc.add_paragraph()  # 创建段落，此处也可直接放入文本，默认字号11
    ph_format = paragraph1.paragraph_format
    ph_format.space_before = Pt(10)  # 设置段前间距
    ph_format.space_after = Pt(10)  # 设置段后间距
    ph_format.line_spacing = Pt(20)  # 设置行间距
    ph_format.first_line_indent = 203200*number_line_indent  # 行开头缩进2字符
    run = paragraph1.add_run(nameStr)  # 创建run，另设字号、粗细
    run.bold = True  # 设置字体为粗体
    run.font.color.rgb = RGBColor(200, 180, 180)
    chg_font(run, fontname='微软雅黑', size=Pt(font_size))  # 设置字体和字号
    return doc


def getDateTimeStr():
    y_str = time.strftime("%Y", time.localtime())
    m_str = time.strftime("%m", time.localtime())
    d_str = time.strftime("%d", time.localtime())
    dateTimeStr = y_str+" 年 "+m_str+" 月 "+d_str+" 日"
    return dateTimeStr


def coverPageMaking(doc, title, subTitle, signature):
    # 标题、空行
    doc = addEmptyRows(doc, number_rows=4, distance=20)
    doc, num_empty_rows = addTitle(doc, title)  #
    doc = addSubTitle(doc, subTitle, font_size=20)
    doc = addEmptyRows(doc, number_rows=num_empty_rows, distance=40)
    # 右下角署名、日期
    doc = addRightStr(doc, signature, number_line_indent=22, font_size=20)
    doc = addRightStr(doc, getDateTimeStr(), number_line_indent=18, font_size=18)
    return doc


def getCoverPage(baseCoverPage_dir, coverPage_dir, title, subTitle, signature):
    if not os.path.exists(baseCoverPage_dir):
        os.makedirs(baseCoverPage_dir)
    if not os.path.exists(coverPage_dir):
        os.makedirs(coverPage_dir)
    coverPage_filepath = os.path.join(baseCoverPage_dir, "coverPage_zyb.docx").replace('\\', '/')
    doc = Document(coverPage_filepath)
    sec = doc.sections[0]  # sections对应文档中的“节”
    # 以下依次设置左、右、上、下页面边距
    distance1 = Cm(2.54)  # 使用“适中”页面布局页边距
    distance2 = Cm(1.91)  # 使用“适中”页面布局页边距
    sec.left_margin = distance2
    sec.right_margin = distance2
    sec.top_margin = distance1
    sec.bottom_margin = distance1
    # 设置页面宽度 高度
    sec.page_width = Cm(20.9)
    sec.page_height = Cm(29.6)

    doc = coverPageMaking(doc, title, subTitle, signature)
    output_filepath = os.path.join(coverPage_dir, "封面.docx").replace('\\', '/')
    doc.save(output_filepath)
    # 封面word转pdf，切图片
    cover_PDF_filepath = output_filepath.replace(".docx", ".pdf")
    doc2PDF(output_filepath, cover_PDF_filepath)
    cover_Img_filepath = output_filepath.replace(".docx", ".png")
    makeCoverImg(cover_PDF_filepath, cover_Img_filepath)
    os.remove(cover_PDF_filepath)
    return output_filepath, doc


from PIL import Image
import fitz


# 重设图片大小
def reset_img(path_filename_in, path_filename_out, height, wide):
    """
    重设图片大小
    :param path_filename_in: 待处理图片带拓展完整路径
    :param path_filename_out: 图片存放带拓展完整路径
    :param height: 高度
    :param wide: 宽度
    """
    im = Image.open(path_filename_in)
    out = im.resize((height, wide), Image.ANTIALIAS)
    im.close()
    out.save(path_filename_out)
    out.close()


# 将第一页作为封面
def makeCoverImg(pdf_path, cover_Img_filepath):
    pdfDoc = fitz.open(pdf_path)  # 转换原pdf
    page = pdfDoc[0]
    pix = page.getPixmap(alpha=False)
    pix.writePNG(cover_Img_filepath)  # 将图片写入指定的文件夹内
    # 设置封面大小 128*128
    reset_img(cover_Img_filepath, cover_Img_filepath, height=128, wide=128)


# ========================================== word页码生成相关 =========================================


# 加页码
def addPageMark(doc_filepath, output_filepath):
    doc = Document(doc_filepath)
    for sec in doc.sections:
        sec.footer_distance = Cm(1.5)
    doc.save(doc_filepath)

    word = client.Dispatch("Word.Application")
    root_path=os.getcwd()
    doc_filepath=root_path+doc_filepath
    doc_filepath=doc_filepath.replace('./','/').replace('\\','/')
    print(doc_filepath)
    WordDoc = word.Documents.Open(doc_filepath)
    try:
        WordDoc.Sections(1).Footers(1).PageNumbers.Add(2, False)  # 页码位置 1 中间 2 居右
        WordDoc.Sections(1).Footers(1).PageNumbers.NumberStyle = 57
    except Exception:
        print("未能成功添加页码")
        print(traceback.print_exc())
    output_filepath = root_path + output_filepath
    output_filepath = output_filepath.replace('./', '/').replace('\\', '/')
    WordDoc.SaveAs(output_filepath)
    print(output_filepath)
    WordDoc.Close()
    word.Quit()


# ============================================= word转pdf =============================================


# from waterMarkPDF.addPageWatermarkPDF import makeWatermarkPDF
#
#
# def doc2PDFwithWatermark(doc_filepath, pdf_filepath):
#     try:
#         word = client.Dispatch("Word.Application")
#         if os.path.exists(pdf_filepath):
#             os.remove(pdf_filepath)
#         worddoc = word.Documents.Open(doc_filepath, ReadOnly=1)
#         worddoc.SaveAs(pdf_filepath, FileFormat=17)
#         worddoc.Close()
#         word.Quit()
#
#         watermark_filepath = 'F:/Muyu/报告创作/PDF转换/waterMark_zyb.jpg'
#         img_pre_dir = 'F:/Muyu/报告创作/PDF转换/img_pre'
#         img_dir = 'F:/Muyu/报告创作/PDF转换/img'
#         makeWatermarkPDF(pdf_filepath, watermark_filepath, pdf_filepath, img_pre_dir, img_dir)
#
#         return pdf_filepath
#     except:
#         return 1


def doc2PDF(doc_filepath, pdf_filepath):
    try:
        word = client.Dispatch("Word.Application")
        if os.path.exists(pdf_filepath):
            os.remove(pdf_filepath)
        worddoc = word.Documents.Open(doc_filepath, ReadOnly=1)
        worddoc.SaveAs(pdf_filepath, FileFormat=17)
        worddoc.Close()
        word.Quit()

        return pdf_filepath
    except:
        return 1

